import pandas as pd
import numpy as np
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


def format_test_paper(num: int):
    ptest = pd.read_excel(f"__cache__\\test_paper{num}.xlsx", sheet_name=0)
    assert isinstance(ptest, pd.DataFrame)
    doc = Document()
    # 设置样式
    from docx.oxml.ns import qn
    doc.styles['Normal'].font.name = u'宋体'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

    doc.add_heading(f"试卷 {num}", level=0)
    doc.add_paragraph("组题: 黄佳俊、温晓平 审稿：李粤", "Subtitle").paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    for ptest_type in ["单选题", "多选题", "填空题"]:
        doc.add_heading(ptest_type, level=2)
        view = ptest[ptest["题型"] == ptest_type].take(np.random.permutation(len(ptest[ptest["题型"] == ptest_type])))[["题干", "选择", "答案", "题型序号"]]
        ptest_num = 0
        for row in view.itertuples():
            ptest_num = ptest_num + 1
            doc.add_paragraph(rf"{ptest_num}、").add_run(rf"【题目】 {row[1]}")
            if isinstance(row[2], str):
                doc.add_paragraph().add_run(row[2])
            doc.add_paragraph(r"【答案】").add_run(str(row[3]))
            doc.add_paragraph(r"【来源】").add_run(rf"Java 题库整理：题型：{ptest_type}, 题型序号: {row[4]}")

    doc.save(f"__cache__\\test_paper{num}.docx")
    # pd.concat([ptest[ptest["题型"] == i].take(np.random.permutation(len(ptest[ptest["题型"] == i])))for i in ["单选题", "多选题", "填空题"]]).to_excel(ptest_path, encoding="utf8")
